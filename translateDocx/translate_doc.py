from docx import Document
import os
import time
from time import sleep
from openai import OpenAI, OpenAIError, APIError, RateLimitError, APIConnectionError, Timeout

DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "")
DEEPSEEK_BASE_URL = os.getenv("DEEPSEEK_BASE_URL", "")  # hoặc "https://api.deepseek.com/v1"
DEEPSEEK_MODEL = "deepseek-chat"  # hoặc "deepseek-reasoner" tùy bạn dùng mode nào

client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url=DEEPSEEK_BASE_URL)

def translate_text(text: str, temperature: float = 0.7) -> str:
    """
    Dịch `text` sang `target_lang` bằng DeepSeek API.
    Chỉ trả về phần dịch, không có lời giải thích nào khác.
    """
    if not text.strip():
        return text

    # Chuẩn bị prompt (system + user)
    messages = [
    {
        "role": "system",
        "content": (
            "You are a professional translator specializing in poker literature and strategy. "
            "Your task is to translate the given English text into natural, fluent, and accurate Vietnamese, "
            "while preserving the original meaning, tone, and technical poker terminology.\n\n"
            "Translation rules:\n"
            "1. Translate exactly — do not add, omit, interpret, or explain anything beyond what is written.\n"
            "2. Keep all poker-related terms (such as range, EV, GTO, equity, bluff, combo, exploitative play, etc.) "
            "either in English or translated to their most common Vietnamese equivalent if one exists naturally.\n"
            "3. Maintain the original paragraph structure and sentence flow.\n"
            "4. Ensure the translation sounds professional, readable, and natural in Vietnamese, not like a machine translation.\n"
            "5. Do not include commentary, analysis, or summaries. Output only the translated text.\n\n"
            "Context conditions to improve translation quality:\n"
            "- The source text comes from a professional poker strategy book for intermediate to advanced players.\n"
            "- Prioritize accuracy in strategic and mathematical terms (e.g., pot odds, equity realization, bet sizing).\n"
            "- Use neutral and formal tone suitable for educational or instructional writing.\n"
            "- If a phrase has no direct equivalent in Vietnamese, keep it in English and translate the surrounding context naturally.\n"
            "- Pay attention to nuances that may affect meaning (e.g., exploitative vs GTO, range advantage, polarized hands)."
        )
    },
    {
        "role": "user",
        "content": (
            "Input:\n" + text
        )
    }
]
    for attempt in range(1, 4):
            try:
                response = client.chat.completions.create(
                    model=DEEPSEEK_MODEL,
                    messages=messages,
                    temperature=temperature,
                    timeout=30,
                    stream=False# giây
                )
                translated = response.choices[0].message.content.strip()
                return translated or text  # nếu API trả rỗng thì fallback về text gốc

            except (APIError, APIConnectionError, Timeout, RateLimitError) as e:
                print(f"[⚠️ API Error: - Attempt {attempt}/{3}] {e}")
                if attempt < 3:
                    time.sleep(2 * attempt)
                    continue
                else:
                    print("[❌ API Error: DeepSeek API failed — returning original text]")
                    return text  # fallback

            except OpenAIError as e:
                print(f"[❌ API Error: SDK Error: {e}]")
                return text  # fallback
            except Exception as e:
                print(f"[💥 API Error: Unexpected Error: {e}]")
                return text  # fallback
    # Nếu hết retries mà vẫn không thành công
    return text

def get_run_style(run):
    """Trích xuất các thuộc tính style của run và trả về dict."""
    return {
        "bold": run.bold,
        "italic": run.italic,
        "underline": run.underline,
        "style": run.style,
        "font_name": run.font.name,
        "font_size": run.font.size,
        "color": run.font.color.rgb if run.font.color else None,
    }

# ----------- HÀM DỊCH FILE DOCX -------------
def translate_docx(input_path, output_path):    
    # Mở tài liệu
    doc = Document(input_path)
    
    # Dịch các paragraph
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            translated_text = translate_text(paragraph.text)
            print(f"Dịch đoạn: {paragraph.text} -> {translated_text}")
            runs_data = []
            for run in paragraph.runs:
                runs_data.append(get_run_style(run=run))
            paragraph.clear()
            new_run = paragraph.add_run(translated_text)
            for item in runs_data:
                new_run.bold = item["bold"]
                new_run.italic = item["italic"]
                new_run.underline = item["underline"]
                new_run.style = item["style"]
                new_run.font.name = item["font_name"]
            sleep(0.1)
    # Lưu tài liệu đã dịch
    doc.save(output_path)
    print(f"Đã dịch xong và lưu tại: {output_path}")

# ----------- CHẠY -------------
if __name__ == "__main__":
    input_file = "/Users/hung/Documents/PythonTraining/PythonTraining/translateDocx/input.docx"     # đường dẫn file gốc
    output_file = "/Users/hung/Documents/PythonTraining/PythonTraining/translateDocx/output_vi.docx"  # file kết quả
    translate_docx(input_file, output_file) 
