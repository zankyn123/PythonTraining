import os
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import List, Tuple, Dict, Any
from docx import Document
import time
from time import sleep
import datetime
import sys
from tqdm import tqdm
from openai import OpenAI, OpenAIError, APIError, RateLimitError, APIConnectionError, Timeout

DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "")
DEEPSEEK_BASE_URL = os.getenv("DEEPSEEK_BASE_URL", "")
DEEPSEEK_MODEL = "deepseek-chat" 
# My macmini has 8 core -> can run 10 threads comfortably
MAX_WORKERS = int(os.getenv("MAX_WORKERS", "10"))
WORD_LIMIT = int(os.getenv("BATCH_WORD_LIMIT", "400"))

def chunk_texts_by_words(texts: List[str], word_limit: int = WORD_LIMIT) -> List[List[str]]:
    """Gom nhóm các đoạn văn thành các batch sao cho tổng số từ của mỗi batch không vượt quá word_limit."""
    
    # Khởi tạo danh sách chứa các batch, batch hiện tại và số từ trong batch hiện tại
    batches, current_batch, current_count = [], [], 0
    
    # Lặp qua từng đoạn văn trong danh sách texts
    for t in texts:
        # Tính số từ của đoạn văn hiện tại bằng cách tách theo khoảng trắng
        wc = len(t.split())
        
        # Nếu đoạn văn có số từ lớn hơn giới hạn và batch hiện tại đang rỗng,
        # nghĩa là không thể gộp với đoạn nào khác, thì thêm đoạn này vào một batch riêng
        if wc > word_limit and not current_batch:
            batches.append([t])
            continue
        
        # Nếu số từ sau khi thêm đoạn hiện tại vượt quá giới hạn và batch hiện tại không rỗng,
        # thì lưu batch hiện tại và bắt đầu một batch mới chứa đoạn này
        if current_count + wc > word_limit and current_batch:
            batches.append(current_batch)
            current_batch, current_count = [t], wc
        else:
            # Nếu không vượt quá giới hạn, thêm đoạn vào batch hiện tại và cập nhật số từ
            current_batch.append(t)
            current_count += wc
    
    # Sau vòng lặp, nếu còn batch không được thêm vào danh sách batches, thêm nó vào
    if current_batch:
        batches.append(current_batch)
    
    return batches


def translate_batch_openai(client: OpenAI, texts: List[str]) -> List[str]:
    """Dịch 1 batch text qua OpenAI API (DeepSeek-compatible nếu có base_url)."""
    outputs = []
    
    
    for t in texts:
        translated = translate_text(client, t)
        outputs.append(translated)
    return outputs

def translate_text(client: OpenAI, text: str, temperature: float = 0.7) -> str:
    """
    Dịch `text` sang `target_lang` bằng DeepSeek API.
    Chỉ trả về phần dịch, không có lời giải thích nào khác.
    """
    if not text.strip():
        return text

    # Chuẩn bị prompt (system + user)
    messages = [
    {
        "role": "system",
        "content": (
            "You are a professional translator specializing in poker literature and strategy. "
            "Your task is to translate the given English text into natural, fluent, and accurate Vietnamese, "
            "while preserving the original meaning, tone, and technical poker terminology.\n\n"
            "Translation rules:\n"
            "1. Translate exactly — do not add, omit, interpret, or explain anything beyond what is written.\n"
            "2. Keep all poker-related terms (such as range, GTO, EV, range, equity, bluff, combo, exploitative play, solver, polarize, depolarize, merge, capped range, uncapped range, pot odds, value bet, thin value, overbet, underbet, c-bet, 3-bet, 4-bet, flat call, limp, squeeze, board texture, equity realization, blocker, nut advantage, range advantage, hand, flop, turn, river, value bet, node blockers, pot multiway, multiway, high card, kicker, suited connectors, pocket pairs, one-gappers, two-gappers, backdoor flush draw, gutshot straight draw, open-ended straight draw, overcards, undercards, top pair, second pair, bottom pair, middle pair, overpair, set, trips, full house, flush, straight, wheel, nut flush, nut straight, cooler, bad beat)"
            "either in English or translated to their most common Vietnamese equivalent if one exists naturally.\n"
            "3. Maintain the original paragraph structure and sentence flow.\n"
            "4. Ensure the translation sounds professional, readable, and natural in Vietnamese, not like a machine translation.\n"
            "5. Do not include commentary, analysis, or summaries. Output only the translated text.\n\n"
            "Context conditions to improve translation quality:\n"
            "- The source text comes from a professional poker strategy book for intermediate to advanced players.\n"
            "- Prioritize accuracy in strategic and mathematical terms (e.g., pot odds, equity realization, bet sizing).\n"
            "- Use neutral and formal tone suitable for educational or instructional writing.\n"
            "- If a phrase has no direct equivalent in Vietnamese, keep it in English and translate the surrounding context naturally.\n"
            "- Pay attention to nuances that may affect meaning (e.g., exploitative vs GTO, range advantage, polarized hands)."
        )
    },
    {
        "role": "user",
        "content": (
            "Input:\n" + text
        )
    }
]
    for attempt in range(1, 4):
            try:
                response = client.chat.completions.create(
                    model=DEEPSEEK_MODEL,
                    messages=messages,
                    temperature=temperature,
                    timeout=30,
                    stream=False# giây
                )
                translated = response.choices[0].message.content.strip()
                return translated or text  # nếu API trả rỗng thì fallback về text gốc

            except (APIError, APIConnectionError, Timeout, RateLimitError) as e:
                print(f"[⚠️ API Error: - Attempt {attempt}/{3}] {e}")
                if attempt < 3:
                    time.sleep(2 * attempt)
                    continue
                else:
                    print("[❌ API Error: DeepSeek API failed — returning original text]")
                    return text  # fallback

            except OpenAIError as e:
                print(f"[❌ API Error: SDK Error: {e}]")
                return text  # fallback
            except Exception as e:
                print(f"[💥 API Error: Unexpected Error: {e}]")
                return text  # fallback
    # Nếu hết retries mà vẫn không thành công
    return text

# ----------------------
# DOCX Processing
# ----------------------
def collect_translatable_items(doc: Document) -> List[Tuple[str, object]]:
    """Thu thập tất cả đoạn văn cần dịch (paragraph + table cell)."""
    items = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            items.append((txt, ("paragraph", p)))
    for table in doc.tables:
        for r in table.rows:
            for c in r.cells:
                cell_text = "\n".join([pp.text for pp in c.paragraphs]).strip()
                if cell_text:
                    items.append((cell_text, ("table-cell", c)))
    return items

def get_run_style(run):
    """Trích xuất các thuộc tính style của run và trả về dict."""
    return {
        "bold": run.bold,
        "italic": run.italic,
        "underline": run.underline,
        "style": run.style,
        "font_name": run.font.name,
        "font_size": run.font.size,
        "color": run.font.color.rgb if run.font.color else None,
    }

def apply_translations(doc: Document, translations_map: dict):
    """Thay thế văn bản bằng bản dịch trong docx (ảnh & bảng giữ nguyên)."""
    for p in doc.paragraphs:
        if p.text and p.text in translations_map:
            new_text = translations_map[p.text]
            runs_data = []
            for run in p.runs:
                runs_data.append(get_run_style(run=run))
            p.clear()
            new_run = p.add_run(new_text)
            for item in runs_data:
                new_run.bold = item["bold"]
                new_run.italic = item["italic"]
                new_run.underline = item["underline"]
                new_run.style = item["style"]
                new_run.font.name = item["font_name"]

    for table in doc.tables:
        for r in table.rows:
            for c in r.cells:
                original = "\n".join([pp.text for pp in c.paragraphs])
                if original and original in translations_map:
                    translated = translations_map[original]
                    if c.paragraphs:
                        first = c.paragraphs[0]
                        first.clear()
                        first.add_run(translated)


# ----------------------
# Main translation
# ----------------------
def translate_docx_file(input_path: str, output_path: str = None, show_progress: bool = True):
    """
    Dịch file docx với hiển thị tiến độ chi tiết.
    
    Args:
        input_path: Đường dẫn file input
        output_path: Đường dẫn file output (mặc định là {tên_file}_vi.docx)
        show_progress: Hiển thị thanh tiến độ (mặc định là True)
    """
    if not DEEPSEEK_API_KEY:
        raise RuntimeError("❌ OPENAI_API_KEY chưa được thiết lập.")

    # Hiển thị thông tin bắt đầu
    start_time = time.time()
    print(f"🚀 Bắt đầu dịch file: {os.path.basename(input_path)}")
    print(f"⚙️ Cấu hình: MAX_WORKERS={MAX_WORKERS}, WORD_LIMIT={WORD_LIMIT}")
    
    client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url=DEEPSEEK_BASE_URL or None)

    if output_path is None:
        base, ext = os.path.splitext(input_path)
        output_path = f"{base}_vi{ext}"

    # Đọc tài liệu
    print(f"📄 Đang đọc tài liệu: {input_path}")
    doc = Document(input_path)
    
    # Thu thập các phần tử cần dịch
    print("🔍 Đang phân tích tài liệu...")
    items = collect_translatable_items(doc)
    if not items:
        print("⚠️ Không có đoạn văn bản nào cần dịch.")
        doc.save(output_path)
        return output_path

    # Loại bỏ trùng lặp
    unique_texts = []
    for txt, _ in items:
        if txt not in unique_texts:
            unique_texts.append(txt)

    total_texts = len(unique_texts)
    total_words = sum(len(text.split()) for text in unique_texts)
    
    print(f"📊 Thống kê:")
    print(f"  - Tổng đoạn cần dịch: {total_texts}")
    print(f"  - Tổng số từ: {total_words}")

    # Chia thành các batch
    batches = chunk_texts_by_words(unique_texts, word_limit=WORD_LIMIT)
    total_batches = len(batches)
    print(f"📦 Đã chia thành {total_batches} batch để xử lý")
    
    # Khởi tạo biến theo dõi tiến độ
    translations_map = {}
    completed_batches = 0
    completed_texts = 0
    
    # Hiển thị thanh tiến độ tổng thể
    if show_progress:
        progress_bar = tqdm(total=total_texts, desc="Tiến độ dịch", unit="đoạn")

    # Sử dụng ThreadPoolExecutor để xử lý đa luồng
    print(f"⚡ Bắt đầu dịch với {MAX_WORKERS} luồng song song...")
    batch_start_time = time.time()
    
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as exe:
        # Tạo dictionary ánh xạ Future objects với các batch
        future_to_batch = {exe.submit(translate_batch_openai, client, b): b for b in batches}
        
        # Lặp qua các Future objects khi hoàn thành
        for fut in as_completed(future_to_batch):
            # Lấy batch tương ứng với Future đã hoàn thành
            batch = future_to_batch[fut]
            batch_size = len(batch)
            
            try:
                # Lấy kết quả dịch
                results = fut.result()
                
                # Cập nhật tiến độ
                completed_batches += 1
                completed_texts += batch_size
                
                # Hiển thị thông tin tiến độ
                batch_time = time.time() - batch_start_time
                batch_words = sum(len(text.split()) for text in batch)
                
                # In thông tin tiến độ
                print(f"✅ Batch {completed_batches}/{total_batches} hoàn thành: {batch_size} đoạn, {batch_words} từ trong {batch_time:.2f}s")
                print(f"   Tiến độ: {completed_texts}/{total_texts} đoạn ({completed_texts/total_texts*100:.1f}%)")
                
                # Cập nhật thanh tiến độ
                if show_progress:
                    progress_bar.update(batch_size)
                
                # Reset thời gian cho batch tiếp theo
                batch_start_time = time.time()
                
            except Exception as e:
                # Xử lý lỗi batch
                print(f"⚠️ Lỗi batch {completed_batches+1}/{total_batches}: {e}")
                print(f"   → Fallback: dịch từng đoạn một...")
                
                results = []
                # Xử lý từng đoạn văn trong batch
                for i, t in enumerate(batch):
                    try:
                        # Thử dịch từng đoạn riêng lẻ
                        res = translate_batch_openai(client, [t])[0]
                        print(f"   ✓ Đoạn {i+1}/{len(batch)} thành công")
                    except Exception as inner_e:
                        # Nếu vẫn lỗi, giữ nguyên văn bản gốc
                        res = t
                        print(f"   ✗ Đoạn {i+1}/{len(batch)} thất bại: {inner_e}")
                    results.append(res)
                    
                    # Cập nhật thanh tiến độ
                    if show_progress:
                        progress_bar.update(1)
                
                # Cập nhật tiến độ sau khi xử lý fallback
                completed_batches += 1
                completed_texts += batch_size
            
            # Cập nhật translations_map với các cặp (văn bản gốc: bản dịch)
            for src, tr in zip(batch, results):
                translations_map[src] = tr
                
            # Hiển thị ước tính thời gian còn lại
            elapsed_time = time.time() - start_time
            estimated_total_time = elapsed_time / (completed_texts / total_texts) if completed_texts > 0 else 0
            remaining_time = max(0, estimated_total_time - elapsed_time)
            
            # Định dạng thời gian còn lại
            remaining_str = str(datetime.timedelta(seconds=int(remaining_time)))
            print(f"⏱️ Ước tính thời gian còn lại: {remaining_str}")
            
    # Đóng thanh tiến độ
    if show_progress:
        progress_bar.close()

    # Áp dụng bản dịch vào tài liệu
    print("🔄 Đang áp dụng bản dịch vào tài liệu...")
    apply_translations(doc, translations_map)
    
    # Lưu tài liệu
    print(f"💾 Đang lưu tài liệu đã dịch: {output_path}")
    doc.save(output_path)
    
    # Hiển thị thống kê cuối cùng
    end_time = time.time()
    total_time = end_time - start_time
    words_per_second = total_words / total_time if total_time > 0 else 0
    
    print("\n📊 Thống kê cuối cùng:")
    print(f"  - Tổng đoạn đã dịch: {total_texts}")
    print(f"  - Tổng số từ: {total_words}")
    print(f"  - Thời gian thực hiện: {total_time:.2f} giây")
    print(f"  - Tốc độ dịch: {words_per_second:.2f} từ/giây")
    print(f"✅ File đã dịch lưu tại: {output_path}")
    
    return output_path


# ----------------------
# CLI
# ----------------------
if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Dịch file .docx qua OpenAI (hoặc DeepSeek endpoint tương thích).")
    parser.add_argument("input", help="Đường dẫn file .docx cần dịch")
    parser.add_argument("-o", "--output", help="Đường dẫn file .docx sau khi dịch (tuỳ chọn)")
    parser.add_argument("--no-progress", action="store_true", help="Không hiển thị thanh tiến độ")
    parser.add_argument("--workers", type=int, help=f"Số luồng xử lý song song (mặc định: {MAX_WORKERS})")
    parser.add_argument("--word-limit", type=int, help=f"Giới hạn số từ mỗi batch (mặc định: {WORD_LIMIT})")
    
    args = parser.parse_args()
    
    # Cập nhật cấu hình nếu có
    if args.workers:
        MAX_WORKERS = args.workers
    if args.word_limit:
        WORD_LIMIT = args.word_limit
    
    # Chạy dịch với cấu hình đã thiết lập
    translate_docx_file(args.input, args.output, show_progress=not args.no_progress)